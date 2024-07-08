##
## Use an OpenAI GPT model on Microsoft to generate a meeting note from a .VTT transcript file
##
## Record Sure Limited ('Recordsure') licences this file to you under the MIT license
##

import os
import webvtt
import re
from tqdm import tqdm
from dotenv import load_dotenv
from openai import AzureOpenAI
from collections import defaultdict
from docx import Document
from datetime import date
import pandas
from argparse import ArgumentParser

parser = ArgumentParser(prog='generateMeetingNote.py')
parser.add_argument('vttPath', nargs='?', help='Path to .vtt file to be used in generation', default=os.path.join("ExampleTranscript", "ExampleTranscript1.vtt"))
in_file = parser.parse_args().vttPath

# The maximum transcript chunking size
max_segment_length = 8190
# A regex used to split transcript sections based on sentence boundary punctuation
rx_sentence_split = re.compile('(?<!etc)(?<!mr)(?<!mrs)(?<!ms)(?<!dr)(?<!\\d)[\\.\\!\\?\u2022](?=\\s+.)', flags=re.UNICODE | re.IGNORECASE)
# Whether or not to include the 'voices' (Speaker labels) in the GPT prompts and therefore outputs
use_speaker_labels = True

# Load the core prompts to be used
core_prompts_df = pandas.read_excel('Prompts.xlsx', sheet_name='Core Prompts')
core_prompts_df.applymap(lambda x: x.strip() if isinstance(x, str) else x)
key_points_summarization_prompt = " ".join(core_prompts_df.query('Key=="key_points_summarization"')['Prompt'].values)
overall_summarization_prompt = " ".join(core_prompts_df.query('Key=="overall_summarization"')['Prompt'].values)
topic_analysis_prelude_prompt = " ".join(core_prompts_df.query('Key=="topic_analysis_prelude"')['Prompt'].values)
topic_analysis_postlude_prompt = " ".join(core_prompts_df.query('Key=="topic_analysis_postlude"')['Prompt'].values)

# Load the topic segmentation prompts to be used
topic_prompts_df = pandas.read_excel('Prompts.xlsx', sheet_name='Topic Prompts')
topic_prompts_df.applymap(lambda x: x.strip() if isinstance(x, str) else x)
topic_list = topic_prompts_df['Topic'].tolist()
rx_topic_search = re.compile("|".join(topic_list), flags=re.UNICODE | re.IGNORECASE)

# Load the Azure OpenAI configuration for your tenancy - typically stored in a .env file in this folder
load_dotenv()
OPENAI_API_TYPE = os.getenv('OPENAI_API_TYPE')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
OPENAI_API_VERSION = os.getenv('OPENAI_API_VERSION')
OPENAI_API_BASE = os.getenv('OPENAI_API_BASE')
OPENAI_API_DEPLOYMENT_NAME = os.getenv('OPENAI_API_DEPLOYMENT_NAME')

# Initialise an instance of an AzureOpenAI client to work with
client = AzureOpenAI(azure_endpoint=OPENAI_API_BASE, azure_deployment=OPENAI_API_DEPLOYMENT_NAME, api_version=OPENAI_API_VERSION, api_key=OPENAI_API_KEY)

# Function to split the VTT into simple text chunks
def split_transcript(vtt, max_length):
    # Split transcript into segments based on max_length
    segments = []
    current_segment = ""
    for caption in tqdm(vtt):
        for sentence in rx_sentence_split.split(caption.text):
            if len(current_segment) + 1 + len(sentence) + ((2 + len(caption.voice)) if use_speaker_labels else 0) <= max_length:
                current_segment += ("" if (sentence == "") else " ") + ((caption.voice + ": ") if use_speaker_labels else "") + sentence
            else:
                segments.append(current_segment)
                current_segment = sentence
    if current_segment:
        segments.append(current_segment)
    return segments

# Function calling the GPT to get the key points from a transcript segment
def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model= OPENAI_API_DEPLOYMENT_NAME,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": key_points_summarization_prompt

            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

# Function calling the GPT to consolidate the summary output from each transcript segment into a single overall summary
def combined_summary_extraction(combined_summary):
    response = client.chat.completions.create(
        model= OPENAI_API_DEPLOYMENT_NAME,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": overall_summarization_prompt
            },
            {
                "role": "user",
                "content": combined_summary
            }
        ]
    )
    return response.choices[0].message.content

# Function calling the GPT model to sort the summarised items into their likely topics 
def perform_zero_shot_classification(text):
    system_prompt = f"""
                    {topic_analysis_prelude_prompt}
                    ```
                    Criteria: Classify conversations under the following categories:
                    """
    for i, topic in topic_prompts_df.iterrows():
        system_prompt += f"""- "{topic['Topic']}" {topic['Description']}
                    """
    system_prompt += f"""```
                    {topic_analysis_postlude_prompt}
                    """
    
    response = client.chat.completions.create(
                model=OPENAI_API_DEPLOYMENT_NAME,
                messages=[ { "role": "system", "content": system_prompt }, { "role": "user", "content": text } ],
                temperature=0
            )

    return response.choices[0].message.content

# Read the VTT input file
vtt = webvtt.read(in_file)
# Split it into right-sized, plain-text chunks
print("Generating transcript segments")
segments = split_transcript(vtt, max_segment_length)

# Generate summary items for each chunk
summaries = []
print("\n\nGenerating topic summaries\n")
for segment in tqdm(segments):
    summary = abstract_summary_extraction(segment)
    summaries.append(summary)

# Glue them into a single text blob
combined_summary = " ".join(summaries)

# Generate the overall summary of summaries
print("\n\nGenerating overall summary\n")
overall_summary = combined_summary_extraction(combined_summary)


# Sort the summary results into their likely topics
print("\n\nCategorising summary bullets by topic\n")
classified_items = defaultdict(list)
for item in tqdm(rx_sentence_split.split(combined_summary)):
    category = perform_zero_shot_classification(item)
    category_match = rx_topic_search.search(category)
    if category_match:
        category = category_match.group()
        classified_items[category].append(item.strip())

# Output a Word document in the same location as the input VTT file
print("\n\nGenerating the meeting note\n")
document = Document()
document.add_heading('Meeting Note {:s}'.format(date.today().strftime('%A %d %b %Y')), 0)
disclaimer = document.add_paragraph('')
disclaimer_format = disclaimer.add_run('This meeting note has been produced using AI and the content should be verified before taking or influencing any action regards the client.')
disclaimer_format.bold = True
disclaimer_format.italic = True
document.add_heading('Conversation summary', level=1)

for item in rx_sentence_split.split(overall_summary):
    document.add_paragraph(item.strip() + '.', style='No Spacing')

for topic in topic_list:
    document.add_heading(f"{topic}", level=2)
    for bullet in classified_items[topic]:
        document.add_paragraph(bullet.strip(), style='List Bullet')

out_dir = os.path.dirname(in_file)
out_file = os.path.join(out_dir, os.path.splitext(os.path.basename(in_file))[0] + "_meeting_note.docx")
document.save(out_file)

print("\n\nDone.")
